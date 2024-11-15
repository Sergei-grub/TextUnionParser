import docx
import re
import os
import json
from collections import Counter
from nltk.stem import SnowballStemmer

from TextUnionParser.Analyzer import MessageBox


class DocumentAnalyzer:
    def __init__(self, file_path, benefits_file):
        self.file_path = file_path
        if not os.path.exists(self.file_path):
            MessageBox.show_warning(f"Файл не найден: {self.file_path}")
            self.text = "УУУ"  # Устанавливаем текст в пустую строку
            return  # Выходим из конструктора, если файл не найден
        self.text = self.read_document()
        self.stemmer = SnowballStemmer("russian")  # Используем стеммер для русского языка
        self.benefits_dict, self.block_keyword = self.load_benefits(benefits_file)  # Загружаем льготы из файла
        self.benefits_count = Counter()

    def load_benefits(self, benefits_file):
        """Загружает льготы из JSON файла с проверкой целостности."""
        try:
            with open(benefits_file, 'r', encoding='utf-8') as file:
                data = json.load(file)

                # Проверяем наличие ключей в загруженных данных
                if "block_keyword" not in data or "benefits" not in data:
                    raise ValueError("Файл не содержит необходимые ключи.")

                block_keyword = data["block_keyword"]
                benefits_dict = data["benefits"]

                # Проверяем, что benefits_dict является словарем
                if not isinstance(benefits_dict, dict):
                    raise ValueError("Раздел 'benefits' должен быть словарем.")

                return benefits_dict, block_keyword

        except (FileNotFoundError, json.JSONDecodeError, ValueError) as e:
            # Выводим предупреждение в случае ошибки
            MessageBox.show_warning(f"Ошибка при загрузке файла: {str(e)}")
            return {}, ""  # Возвращаем пустые значения в случае ошибки

    def read_document(self):
        """Чтение текста из документа Word."""
        try:
            doc = docx.Document(self.file_path)
        except Exception as e:
            MessageBox.show_warning(
                f"Ошибка при открытии документа. Проверьте целостность документа и наличия в нем текста. {e}")
            return ""  # Возвращаем пустую строку в случае ошибки

        # Проверяем, есть ли параграфы в документе
        if not doc.paragraphs:
            MessageBox.show_warning("Документ пуст. Пожалуйста, проверьте содержимое.")
            return ""  # Возвращаем пустую строку, чтобы избежать дальнейших ошибок

        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return ' '.join(text)

    def simple_tokenize(self, text):
        return re.findall(r'\b\w+\b', text.lower())

    def analyze_benefits(self):
        """Анализ текста на наличие уникальных льгот и их подсчет."""
        blocks = self.text.split("Коллективный договор")
        self.block_count = len(blocks) - 1

        # Создаем словарь стеммированных ключевых слов
        stemmed_benefits_dict = {category: [self.stemmer.stem(benefit) for benefit in benefits]
                                 for category, benefits in self.benefits_dict.items()}

        for i in range(1, len(blocks)):
            block = blocks[i]
            words = self.simple_tokenize(block)  # Токенизируем текст

            # Используем множество для хранения найденных льгот
            found_benefits = set()

            # Проверяем наличие фраз
            for category, benefits in self.benefits_dict.items():
                for benefit in benefits:
                    if benefit in block.lower() and benefit not in found_benefits:  # Проверяем наличие фразы в тексте
                        self.benefits_count[(category, benefit)] += 1
                        found_benefits.add(benefit)  # Добавляем в множество

            # Проверяем стеммированные слова
            for category, stemmed_benefits in stemmed_benefits_dict.items():
                for benefit in stemmed_benefits:
                    if benefit in words and benefit not in found_benefits:
                        self.benefits_count[(category, benefit)] += 1
                        found_benefits.add(benefit)  # Добавляем в множество

    def create_report(self, output_file):
        """Создание отчета о льготах и преференциях."""
        doc = docx.Document()

        # Добавляем информацию о количестве блоков
        doc.add_heading(f"Коллективный договор: {self.block_count}", level=1)

        unique_categories = set(cat[0] for cat in self.benefits_count.keys())

        for category in unique_categories:
            doc.add_heading(category, level=1)
            for benefit in self.benefits_count:
                if benefit[0] == category:
                    doc.add_paragraph(f"{benefit[1]}: {self.benefits_count[benefit]}")

        try:
            doc.save(output_file)
            # Переместите вывод сообщения об успешном сохранении сюда
            MessageBox.show_info(f"Отчет (report.docx) успешно сохранен в папку {output_file}.")
        except PermissionError:
            MessageBox.show_warning(f"Ошибка: файл '{output_file}' открыт. Закройте файл и попробуйте снова.")
        except Exception as e:
            MessageBox.show_warning(f"Произошла ошибка при сохранении отчета: {e}")
